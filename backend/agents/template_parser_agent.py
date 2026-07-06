"""Agent 3 — Template Parser.

Analyses a template *file* (PDF / DOCX / image) and produces a structured
JSON description: layout, colors, fonts, sections, spacing, columns — plus a
preview image. Used for user-uploaded templates and any downloaded files.
"""
from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from models.schemas import (
    AgentResult,
    TemplateColors,
    TemplateFonts,
    TemplateLayout,
    TemplateMeta,
)
from services.preview_service import generate_file_preview
from .base_agent import BaseAgent
from .resume_reader_agent import SECTION_ALIASES

_HEX = "#{:02x}{:02x}{:02x}"


class TemplateParserAgent(BaseAgent):
    name = "template_parser"

    # ---------------------------------------------------------------- parse

    def parse(self, path: Path) -> tuple[TemplateLayout, TemplateColors, TemplateFonts, list[str]]:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(path)
        if ext == ".docx":
            return self._parse_docx(path)
        if ext in (".png", ".jpg", ".jpeg"):
            return self._parse_image(path)
        return TemplateLayout(), TemplateColors(), TemplateFonts(), []

    def _parse_pdf(self, path: Path):
        import fitz

        layout = TemplateLayout()
        colors = TemplateColors()
        fonts = TemplateFonts()
        sections: list[str] = []
        with fitz.open(path) as doc:
            if doc.page_count == 0:
                return layout, colors, fonts, sections
            page = doc[0]
            width = page.rect.width

            # columns: cluster text blocks by x position
            blocks = [b for b in page.get_text("blocks") if b[6] == 0]
            left = sum(1 for b in blocks if b[0] < width * 0.45)
            right = sum(1 for b in blocks if b[0] >= width * 0.45)
            if left and right and min(left, right) >= max(2, len(blocks) // 5):
                layout.columns = 2
                layout.sidebar = "left" if left < right else "right"

            # fonts and sizes from spans
            sizes: list[float] = []
            font_names: Counter[str] = Counter()
            text_all: list[str] = []
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        sizes.append(round(span["size"], 1))
                        font_names[span["font"]] += 1
                        text_all.append(span["text"])
            if sizes:
                sizes.sort()
                fonts.base_size_pt = sizes[len(sizes) // 2]
                fonts.name_size_pt = sizes[-1]
            if font_names:
                fonts.body = font_names.most_common(1)[0][0]
                fonts.heading = fonts.body

            # colors: render page and take dominant non-background colors
            colors = self._dominant_colors_from_pixmap(page)

            # sections present in the text
            joined = "\n".join(text_all).lower()
            sections = self._detect_sections(joined)

            # spacing: ratio of text area to page area
            text_area = sum((b[2] - b[0]) * (b[3] - b[1]) for b in blocks)
            density = text_area / (page.rect.width * page.rect.height + 1)
            layout.spacing = "compact" if density > 0.55 else ("relaxed" if density < 0.3 else "normal")

            # header style: is the top block centered?
            if blocks:
                top = min(blocks, key=lambda b: b[1])
                center = (top[0] + top[2]) / 2
                layout.header_style = "centered" if abs(center - width / 2) < width * 0.12 else "left"
        return layout, colors, fonts, sections

    def _dominant_colors_from_pixmap(self, page) -> TemplateColors:
        import io

        from PIL import Image

        pix = page.get_pixmap(dpi=50)
        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        img = img.resize((64, 82))
        counts = Counter(img.getdata())
        ranked = [c for c, _ in counts.most_common(30)]
        background = ranked[0] if ranked else (255, 255, 255)

        def distance(a, b) -> int:
            return sum(abs(x - y) for x, y in zip(a, b))

        accents = [c for c in ranked[1:] if distance(c, background) > 120]
        text_color = accents[0] if accents else (34, 34, 34)
        accent = next(
            (c for c in accents[1:] if max(c) - min(c) > 40),  # something saturated
            text_color,
        )
        return TemplateColors(
            background=_HEX.format(*background),
            text=_HEX.format(*text_color),
            primary=_HEX.format(*text_color),
            accent=_HEX.format(*accent),
        )

    def _parse_docx(self, path: Path):
        import docx

        layout = TemplateLayout()
        colors = TemplateColors()
        fonts = TemplateFonts()
        document = docx.Document(str(path))
        font_names: Counter[str] = Counter()
        sizes: list[float] = []
        texts: list[str] = []
        for para in document.paragraphs:
            texts.append(para.text)
            for run in para.runs:
                if run.font.name:
                    font_names[run.font.name] += 1
                if run.font.size is not None:
                    sizes.append(run.font.size.pt)
                if run.font.color and run.font.color.rgb:
                    rgb = str(run.font.color.rgb)
                    if rgb not in ("000000", "FFFFFF") and colors.accent == TemplateColors().accent:
                        colors.accent = f"#{rgb.lower()}"
        if font_names:
            fonts.body = font_names.most_common(1)[0][0]
            fonts.heading = fonts.body
        if sizes:
            fonts.base_size_pt = sorted(sizes)[len(sizes) // 2]
            fonts.name_size_pt = max(sizes)
        if document.tables:
            layout.columns = 2
        sections = self._detect_sections("\n".join(texts).lower())
        return layout, colors, fonts, sections

    def _parse_image(self, path: Path):
        from PIL import Image

        layout = TemplateLayout()
        fonts = TemplateFonts()
        with Image.open(path) as img:
            rgb = img.convert("RGB").resize((64, 82))
            counts = Counter(rgb.getdata())
            ranked = [c for c, _ in counts.most_common(30)]
        background = ranked[0] if ranked else (255, 255, 255)
        distinct = [c for c in ranked[1:] if sum(abs(x - y) for x, y in zip(c, background)) > 120]
        text_color = distinct[0] if distinct else (34, 34, 34)
        accent = next((c for c in distinct[1:] if max(c) - min(c) > 40), text_color)
        colors = TemplateColors(
            background=_HEX.format(*background),
            text=_HEX.format(*text_color),
            primary=_HEX.format(*text_color),
            accent=_HEX.format(*accent),
        )
        sections: list[str] = []
        try:
            import pytesseract

            with Image.open(path) as img:
                text = pytesseract.image_to_string(img).lower()
            sections = self._detect_sections(text)
        except Exception:  # noqa: BLE001 - OCR is optional for template analysis
            self.logger.info("OCR unavailable; skipping section detection for image template")
        return layout, colors, fonts, sections

    def _detect_sections(self, text: str) -> list[str]:
        found: list[str] = []
        for canonical, aliases in SECTION_ALIASES.items():
            if any(re.search(rf"\b{re.escape(a)}\b", text) for a in aliases):
                found.append(canonical)
        return found

    # ------------------------------------------------------------------ run

    def run(self, path: Path, meta: TemplateMeta) -> AgentResult:
        """Analyse the file and enrich the given TemplateMeta in place."""
        try:
            layout, colors, fonts, sections = self.parse(path)
        except Exception as exc:  # noqa: BLE001
            return self.fail(f"Template analysis failed: {exc}")
        meta.layout = layout
        meta.colors = colors
        meta.fonts = fonts
        meta.sections = sections or meta.sections
        preview = Path(path).parent / "previews" / f"{meta.id}.png"
        result = generate_file_preview(path, preview)
        if result is not None:
            meta.preview_path = str(result)
        return self.ok(
            "Template analysed",
            template=meta.model_dump(mode="json"),
        )
