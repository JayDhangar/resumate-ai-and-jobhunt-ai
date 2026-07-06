"""Agent 1 — Resume Reader.

Accepts PDF / DOCX / PNG / JPG / JPEG uploads, extracts raw text (PDF parser,
DOCX parser, or OCR for images), then structures it into :class:`ResumeData`.
Structuring uses the configured LLM when available and falls back to a
deterministic heuristic parser so the pipeline never dead-ends.
"""
from __future__ import annotations

import re
from pathlib import Path

from core.exceptions import ExtractionError, LLMUnavailableError
from models.schemas import (
    ContactLinks,
    EducationItem,
    ExperienceItem,
    ResumeData,
    SkillGroup,
)
from .base_agent import BaseAgent

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}")
URL_RE = re.compile(r"(?:https?://|www\.)[^\s|,;)]+", re.IGNORECASE)
DATE_RANGE_RE = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4})"
    r"\s*[-–—to]+\s*"
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4}|Present|Current)",
    re.IGNORECASE,
)

SECTION_ALIASES: dict[str, list[str]] = {
    "summary": ["summary", "professional summary", "profile", "objective", "about", "about me"],
    "skills": ["skills", "technical skills", "core competencies", "technologies", "tech stack"],
    "experience": ["experience", "work experience", "employment", "professional experience", "work history"],
    "education": ["education", "academics", "academic background", "qualifications"],
    "projects": ["projects", "personal projects", "key projects", "portfolio"],
    "certifications": ["certifications", "certificates", "licenses"],
    "languages": ["languages", "language proficiency"],
    "awards": ["awards", "honors", "achievements", "accomplishments"],
}

STRUCTURING_SYSTEM_PROMPT = """You are a resume parsing expert. Convert raw resume text into JSON with this exact shape:
{
  "name": str, "headline": str, "email": str, "phone": str, "location": str, "summary": str,
  "links": {"linkedin": str, "github": str, "website": str, "other": [str]},
  "skills": [{"category": str, "skills": [str]}],
  "experience": [{"title": str, "company": str, "location": str, "start_date": str, "end_date": str, "current": bool, "bullets": [str]}],
  "education": [{"degree": str, "institution": str, "location": str, "start_date": str, "end_date": str, "gpa": str, "details": [str]}],
  "projects": [{"name": str, "description": str, "technologies": [str], "link": str, "bullets": [str]}],
  "certifications": [{"name": str, "issuer": str, "date": str, "link": str}],
  "languages": [{"name": str, "proficiency": str}],
  "awards": [{"title": str, "issuer": str, "date": str, "description": str}]
}
Extract EVERYTHING present. Never invent information. Use "" or [] for anything absent."""


class ResumeReaderAgent(BaseAgent):
    name = "resume_reader"

    # ------------------------------------------------------------------ text

    def extract_text(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext == ".pdf":
            text = self._extract_pdf(path)
        elif ext == ".docx":
            text = self._extract_docx(path)
        elif ext in (".png", ".jpg", ".jpeg"):
            text = self._extract_image_ocr(path)
        else:
            raise ExtractionError(f"Unsupported file type: {ext}")
        text = text.strip()
        if not text:
            raise ExtractionError(
                "No text could be extracted. If this is a scanned document, "
                "install Tesseract OCR (https://github.com/UB-Mannheim/tesseract/wiki)."
            )
        return text

    def _extract_pdf(self, path: Path) -> str:
        import fitz  # PyMuPDF

        with fitz.open(path) as doc:
            text = "\n".join(page.get_text("text") for page in doc)
            if text.strip():
                return text
            # scanned PDF -> rasterise pages, then AI vision (or Tesseract) reads them
            self.logger.info("PDF has no text layer; using vision/OCR")
            parts: list[str] = []
            for page in doc:
                pix = page.get_pixmap(dpi=200)
                parts.append(self._ocr_bytes(pix.tobytes("png")))
            return "\n".join(parts)

    def _extract_docx(self, path: Path) -> str:
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    def _extract_image_ocr(self, path: Path) -> str:
        mime = "image/jpeg" if path.suffix.lower() in (".jpg", ".jpeg") else "image/png"
        return self._ocr_bytes(path.read_bytes(), mime=mime)

    def _ocr_bytes(self, image_bytes: bytes, mime: str = "image/png") -> str:
        """Read text from an image: AI vision first (no install, better accuracy),
        Tesseract as the offline fallback."""
        if self.llm.available:
            try:
                return self._vision_extract(image_bytes, mime)
            except Exception as exc:  # noqa: BLE001
                self.logger.warning("Vision extraction failed (%s); trying Tesseract", exc)
        return self._tesseract_extract(image_bytes)

    def _vision_extract(self, image_bytes: bytes, mime: str = "image/png") -> str:
        import base64

        return self.llm.complete_vision(
            "You transcribe resume documents. Output the resume's full text content, "
            "preserving section headings and line structure. Output ONLY the transcribed text.",
            "Transcribe this resume image completely.",
            base64.b64encode(image_bytes).decode(),
            mime=mime,
            max_tokens=4000,
            op="vision_resume_read",
        )

    def _tesseract_extract(self, image_bytes: bytes) -> str:
        import io

        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:
            raise ExtractionError(
                "OCR requires pytesseract and Pillow: pip install pytesseract pillow"
            ) from exc
        try:
            image = Image.open(io.BytesIO(image_bytes))
            return pytesseract.image_to_string(image)
        except pytesseract.TesseractNotFoundError as exc:
            raise ExtractionError(
                "Image reading needs either an OpenAI key (AI vision) or the Tesseract binary "
                "(https://github.com/UB-Mannheim/tesseract/wiki)."
            ) from exc

    # ------------------------------------------------------------- structure

    def structure(self, raw_text: str) -> ResumeData:
        """Turn raw text into ResumeData via LLM, falling back to heuristics."""
        if self.llm.available:
            try:
                payload = self.llm.complete_json(
                    STRUCTURING_SYSTEM_PROMPT,
                    f"Resume text:\n\n{raw_text[:15000]}",
                    max_tokens=8000,
                    op="extract_resume",
                )
                return ResumeData.model_validate(payload)
            except (LLMUnavailableError, Exception) as exc:  # noqa: BLE001
                self.logger.warning("LLM structuring failed (%s); using heuristics", exc)
        return self.heuristic_structure(raw_text)

    def heuristic_structure(self, raw_text: str) -> ResumeData:
        """Deterministic section-based parser used when no LLM is available."""
        lines = [ln.rstrip() for ln in raw_text.splitlines()]
        resume = ResumeData()

        if match := EMAIL_RE.search(raw_text):
            resume.email = match.group(0)
        if match := PHONE_RE.search(raw_text):
            resume.phone = match.group(0).strip()
        for url in URL_RE.findall(raw_text):
            low = url.lower()
            if "linkedin" in low and not resume.links.linkedin:
                resume.links.linkedin = url
            elif "github" in low and not resume.links.github:
                resume.links.github = url
            elif not resume.links.website:
                resume.links.website = url
            else:
                resume.links.other.append(url)

        # name: first short line without contact info
        for line in lines[:8]:
            stripped = line.strip()
            if not stripped or EMAIL_RE.search(stripped) or PHONE_RE.search(stripped):
                continue
            if len(stripped.split()) <= 5 and not URL_RE.search(stripped):
                resume.name = stripped
                break

        sections = self._split_sections(lines)
        resume.summary = " ".join(sections.get("summary", [])).strip()
        resume.skills = self._parse_skills(sections.get("skills", []))
        resume.experience = self._parse_experience(sections.get("experience", []))
        resume.education = self._parse_education(sections.get("education", []))
        for key in ("projects", "certifications", "languages", "awards"):
            body = [ln for ln in sections.get(key, []) if ln.strip()]
            if body:
                resume.extra_sections[key] = body
        return resume

    def _split_sections(self, lines: list[str]) -> dict[str, list[str]]:
        alias_lookup = {
            alias: canonical
            for canonical, aliases in SECTION_ALIASES.items()
            for alias in aliases
        }
        sections: dict[str, list[str]] = {}
        current: str | None = None
        for line in lines:
            key = re.sub(r"[^a-z ]", "", line.strip().lower()).strip()
            if key in alias_lookup and len(line.strip()) < 40:
                current = alias_lookup[key]
                sections.setdefault(current, [])
                continue
            if current:
                sections[current].append(line)
        return sections

    def _parse_skills(self, lines: list[str]) -> list[SkillGroup]:
        groups: list[SkillGroup] = []
        for line in lines:
            stripped = line.strip().lstrip("•-*· ")
            if not stripped:
                continue
            if ":" in stripped:
                category, _, rest = stripped.partition(":")
                skills = [s.strip() for s in re.split(r"[,;|]", rest) if s.strip()]
                groups.append(SkillGroup(category=category.strip(), skills=skills))
            else:
                skills = [s.strip() for s in re.split(r"[,;|]", stripped) if s.strip()]
                if skills:
                    groups.append(SkillGroup(category="", skills=skills))
        return groups

    def _parse_experience(self, lines: list[str]) -> list[ExperienceItem]:
        items: list[ExperienceItem] = []
        current: ExperienceItem | None = None
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith(("•", "-", "*", "·")):
                if current is not None:
                    current.bullets.append(stripped.lstrip("•-*· ").strip())
                continue
            if date_match := DATE_RANGE_RE.search(stripped):
                header = DATE_RANGE_RE.sub("", stripped).strip(" |,–—-")
                current = ExperienceItem(
                    start_date=date_match.group(1),
                    end_date=date_match.group(2),
                    current=date_match.group(2).lower() in ("present", "current"),
                )
                parts = re.split(r"\s+[-–—|@]\s+|,\s+", header, maxsplit=1)
                current.title = parts[0].strip()
                if len(parts) > 1:
                    current.company = parts[1].strip()
                items.append(current)
            elif current is None:
                current = ExperienceItem(title=stripped)
                items.append(current)
            elif not current.company:
                current.company = stripped
            else:
                current.bullets.append(stripped)
        return items

    def _parse_education(self, lines: list[str]) -> list[EducationItem]:
        items: list[EducationItem] = []
        current: EducationItem | None = None
        for line in lines:
            stripped = line.strip().lstrip("•-*· ")
            if not stripped:
                continue
            if date_match := DATE_RANGE_RE.search(stripped):
                header = DATE_RANGE_RE.sub("", stripped).strip(" |,–—-")
                current = EducationItem(
                    degree=header,
                    start_date=date_match.group(1),
                    end_date=date_match.group(2),
                )
                items.append(current)
            elif current is None:
                current = EducationItem(degree=stripped)
                items.append(current)
            elif not current.institution:
                current.institution = stripped
            else:
                current.details.append(stripped)
        return items

    # ------------------------------------------------------------------ run

    def run(self, path: Path) -> tuple[ResumeData, str]:
        """Full pipeline: file -> raw text -> structured ResumeData."""
        self.logger.info("Reading resume from %s", path.name)
        raw_text = self.extract_text(path)
        resume = self.structure(raw_text)
        return resume, raw_text
