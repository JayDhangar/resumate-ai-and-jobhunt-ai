"""Agent 7 — Export.

Converts a generated resume into every supported download format:
PDF, DOCX, HTML, PNG, Markdown and JSON. All conversions are pure-Python and
offline: HTML -> PDF via xhtml2pdf, PDF -> PNG via PyMuPDF, DOCX built
programmatically with python-docx.
"""
from __future__ import annotations

import io
import json
from pathlib import Path

from core.config import get_settings
from core.exceptions import ExportError
from models.schemas import AgentResult, ResumeData, TemplateMeta
from .base_agent import BaseAgent

SUPPORTED_FORMATS = ("html", "pdf", "docx", "png", "md", "json")


class ExportAgent(BaseAgent):
    name = "export"

    def __init__(self, **kwargs) -> None:
        super().__init__(**kwargs)
        self.settings = get_settings()

    # -------------------------------------------------------------- formats

    def to_pdf(self, html: str, dest: Path, fit_one_page: bool = False) -> Path:
        """Prefer Chromium print-to-PDF (pixel-perfect, matches the browser
        preview exactly); fall back to xhtml2pdf when Playwright is unavailable."""
        try:
            return self._to_pdf_chromium(html, dest, fit_one_page)
        except Exception as exc:  # noqa: BLE001
            self.logger.warning("Chromium PDF failed (%s); falling back to xhtml2pdf", exc)
        return self._to_pdf_xhtml2pdf(html, dest)

    def _to_pdf_chromium(self, html: str, dest: Path, fit_one_page: bool = False) -> Path:
        from playwright.sync_api import sync_playwright

        # Full-bleed printing: page spacing lives INSIDE the document (body padding),
        # so colored template backgrounds cover the entire page with no white frame.
        a4_width_px = 21.0 / 2.54 * 96    # ≈ 794
        a4_height_px = 29.7 / 2.54 * 96   # ≈ 1122
        with sync_playwright() as p:
            browser = p.chromium.launch()
            try:
                page = browser.new_page()
                page.set_viewport_size({"width": int(a4_width_px), "height": 1100})
                page.set_content(html, wait_until="load")
                scale = 1.0
                if fit_one_page:
                    # pdf(scale=s) lays content out at width/s, THEN shrinks by s —
                    # re-measure at the scaled layout width until it fits (3% safety)
                    for _ in range(4):
                        page.set_viewport_size({"width": int(a4_width_px / scale), "height": 1100})
                        content_height = page.evaluate("document.body.scrollHeight")
                        if content_height * scale <= a4_height_px * 0.99:
                            break
                        scale = max(0.6, round((a4_height_px * 0.97) / content_height, 3))
                    if scale < 1.0:
                        self.logger.info("1-page fit: scaling PDF to %.0f%%", scale * 100)
                page.pdf(
                    path=str(dest),
                    format="A4",
                    print_background=True,
                    scale=scale,
                    margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
                )
            finally:
                browser.close()
        return dest

    def _to_pdf_xhtml2pdf(self, html: str, dest: Path) -> Path:
        from xhtml2pdf import pisa

        buffer = io.BytesIO()
        result = pisa.CreatePDF(io.StringIO(html), dest=buffer, encoding="utf-8")
        if result.err:
            raise ExportError("PDF conversion failed — the template HTML could not be rendered")
        dest.write_bytes(buffer.getvalue())
        return dest

    def to_png(self, pdf_path: Path, dest: Path) -> Path:
        import fitz

        with fitz.open(pdf_path) as doc:
            if doc.page_count == 0:
                raise ExportError("Generated PDF has no pages")
            pix = doc[0].get_pixmap(dpi=150)
            pix.save(str(dest))
        return dest

    def to_docx(self, resume: ResumeData, template: TemplateMeta, dest: Path) -> Path:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        def hex_to_rgb(value: str) -> RGBColor:
            value = (value or "#222222").lstrip("#")
            if len(value) != 6:
                value = "222222"
            return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))

        primary = hex_to_rgb(template.colors.primary)
        accent = hex_to_rgb(template.colors.accent)
        base_pt = Pt(template.fonts.base_size_pt)
        document = docx.Document()

        name_para = document.add_paragraph()
        if template.layout.header_style == "centered":
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(resume.name or "Your Name")
        run.bold = True
        run.font.size = Pt(template.fonts.name_size_pt)
        run.font.color.rgb = primary

        contact_bits = [x for x in (resume.email, resume.phone, resume.location,
                                    resume.links.linkedin, resume.links.github,
                                    resume.links.website) if x]
        if contact_bits:
            contact = document.add_paragraph()
            if template.layout.header_style == "centered":
                contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            crun = contact.add_run("  •  ".join(contact_bits))
            crun.font.size = Pt(template.fonts.base_size_pt - 1)

        def heading(text: str) -> None:
            para = document.add_paragraph()
            run = para.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(template.fonts.section_size_pt)
            run.font.color.rgb = accent

        def body(text: str, bullet: bool = False, bold_prefix: str = "") -> None:
            para = document.add_paragraph(style="List Bullet" if bullet else None)
            if bold_prefix:
                brun = para.add_run(bold_prefix)
                brun.bold = True
                brun.font.size = base_pt
            run = para.add_run(text)
            run.font.size = base_pt

        titles = {"summary": "Summary", "skills": "Skills", "experience": "Experience",
                  "projects": "Projects", "education": "Education",
                  "certifications": "Certifications", "awards": "Awards", "languages": "Languages"}
        titles.update({k.lower(): v for k, v in resume.section_titles.items()})

        for key in (resume.section_order or list(titles)):
            if key == "summary" and resume.summary:
                heading(titles["summary"])
                body(resume.summary)
            elif key == "skills" and resume.skills:
                heading(titles["skills"])
                for group in resume.skills:
                    prefix = f"{group.category}: " if group.category else ""
                    body(", ".join(group.skills), bold_prefix=prefix)
            elif key == "experience" and resume.experience:
                heading(titles["experience"])
                for exp in resume.experience:
                    dates = f"  ({exp.start_date} – {exp.end_date})" if exp.start_date else ""
                    body(f"{exp.company}{' · ' + exp.location if exp.location else ''}{dates}",
                         bold_prefix=exp.title + " — " if exp.title else "")
                    for b in exp.bullets:
                        body(b, bullet=True)
            elif key == "projects" and resume.projects:
                heading(titles["projects"])
                for proj in resume.projects:
                    desc = proj.description or ", ".join(proj.technologies)
                    body(desc, bold_prefix=proj.name + " — " if proj.name else "")
                    for b in proj.bullets:
                        body(b, bullet=True)
            elif key == "education" and resume.education:
                heading(titles["education"])
                for edu in resume.education:
                    dates = f"  ({edu.start_date} – {edu.end_date})" if edu.start_date else ""
                    body(f"{edu.institution}{dates}", bold_prefix=edu.degree + " — " if edu.degree else "")
                    for d in edu.details:
                        body(d, bullet=True)
            elif key == "certifications" and resume.certifications:
                heading(titles["certifications"])
                for cert in resume.certifications:
                    body(f"{cert.name}{' — ' + cert.issuer if cert.issuer else ''}"
                         f"{' (' + cert.date + ')' if cert.date else ''}", bullet=True)
            elif key == "languages" and resume.languages:
                heading(titles["languages"])
                body(", ".join(
                    f"{l.name} ({l.proficiency})" if l.proficiency else l.name
                    for l in resume.languages
                ))
            elif key == "awards" and resume.awards:
                heading(titles["awards"])
                for award in resume.awards:
                    body(f"{award.title}{' — ' + award.issuer if award.issuer else ''}", bullet=True)

        for key, lines in resume.extra_sections.items():
            if lines:
                heading(titles.get(key, key.title()))
                for line in lines:
                    body(line, bullet=True)

        document.save(str(dest))
        return dest

    def to_markdown(self, resume: ResumeData) -> str:
        lines: list[str] = [f"# {resume.name or 'Your Name'}"]
        if resume.headline:
            lines.append(f"*{resume.headline}*")
        contact = " · ".join(x for x in (resume.email, resume.phone, resume.location) if x)
        if contact:
            lines.append(contact)
        links = " · ".join(x for x in (resume.links.linkedin, resume.links.github, resume.links.website) if x)
        if links:
            lines.append(links)
        if resume.summary:
            lines += ["", "## Summary", resume.summary]
        if resume.skills:
            lines += ["", "## Skills"]
            for group in resume.skills:
                prefix = f"**{group.category}:** " if group.category else ""
                lines.append(f"- {prefix}{', '.join(group.skills)}")
        if resume.experience:
            lines += ["", "## Experience"]
            for exp in resume.experience:
                dates = f" ({exp.start_date} – {exp.end_date})" if exp.start_date else ""
                lines.append(f"\n### {exp.title} — {exp.company}{dates}")
                lines += [f"- {b}" for b in exp.bullets]
        if resume.projects:
            lines += ["", "## Projects"]
            for proj in resume.projects:
                lines.append(f"\n### {proj.name}")
                if proj.description:
                    lines.append(proj.description)
                if proj.technologies:
                    lines.append(f"*{', '.join(proj.technologies)}*")
                lines += [f"- {b}" for b in proj.bullets]
        if resume.education:
            lines += ["", "## Education"]
            for edu in resume.education:
                dates = f" ({edu.start_date} – {edu.end_date})" if edu.start_date else ""
                lines.append(f"- **{edu.degree}** — {edu.institution}{dates}")
        if resume.certifications:
            lines += ["", "## Certifications"]
            lines += [f"- {c.name}{' — ' + c.issuer if c.issuer else ''}" for c in resume.certifications]
        if resume.languages:
            lines += ["", "## Languages",
                      ", ".join(f"{l.name} ({l.proficiency})" if l.proficiency else l.name
                                for l in resume.languages)]
        if resume.awards:
            lines += ["", "## Awards"]
            lines += [f"- {a.title}{' — ' + a.issuer if a.issuer else ''}" for a in resume.awards]
        for key, extra in resume.extra_sections.items():
            if extra:
                lines += ["", f"## {key.title()}"]
                lines += [f"- {line}" for line in extra]
        return "\n".join(lines) + "\n"

    # ------------------------------------------------------------------ run

    def run(
        self,
        resume: ResumeData,
        template: TemplateMeta,
        html: str,
        formats: list[str],
        output_name: str = "resume",
    ) -> AgentResult:
        out_dir = Path(self.settings.generated_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        produced: dict[str, str] = {}
        errors: dict[str, str] = {}

        requested = [f.lower().lstrip(".") for f in formats] or ["html"]
        requested = ["png" if f in ("image", "jpg", "jpeg") else f for f in requested]
        if "png" in requested and "pdf" not in requested:
            requested.append("pdf")  # PNG is rendered from the PDF

        for fmt in dict.fromkeys(requested):  # preserve order, dedupe
            try:
                if fmt == "html":
                    path = out_dir / f"{output_name}.html"
                    path.write_text(html, encoding="utf-8")
                    produced["html"] = str(path)
                elif fmt == "pdf":
                    produced["pdf"] = str(self.to_pdf(
                        html, out_dir / f"{output_name}.pdf",
                        fit_one_page=template.layout.page_mode == "one",
                    ))
                elif fmt == "docx":
                    produced["docx"] = str(self.to_docx(resume, template, out_dir / f"{output_name}.docx"))
                elif fmt == "png":
                    pdf_path = Path(produced.get("pdf", ""))
                    if not pdf_path.is_file():
                        pdf_path = self.to_pdf(html, out_dir / f"{output_name}.pdf",
                                               fit_one_page=template.layout.page_mode == "one")
                        produced["pdf"] = str(pdf_path)
                    produced["png"] = str(self.to_png(pdf_path, out_dir / f"{output_name}.png"))
                elif fmt in ("md", "markdown"):
                    path = out_dir / f"{output_name}.md"
                    path.write_text(self.to_markdown(resume), encoding="utf-8")
                    produced["md"] = str(path)
                elif fmt == "json":
                    path = out_dir / f"{output_name}.json"
                    path.write_text(
                        json.dumps(resume.model_dump(mode="json"), indent=2, ensure_ascii=False),
                        encoding="utf-8",
                    )
                    produced["json"] = str(path)
                else:
                    errors[fmt] = f"Unknown format '{fmt}'. Supported: {', '.join(SUPPORTED_FORMATS)}"
            except Exception as exc:  # noqa: BLE001 - report per-format, never abort the batch
                self.logger.exception("Export to %s failed", fmt)
                errors[fmt] = str(exc)

        if not produced:
            return self.fail("All exports failed", errors=errors)
        return self.ok(
            f"Exported {', '.join(produced)}",
            files=produced,
            errors=errors,
        )
